"""Knowledge Base management endpoints.

Full KB lifecycle: upload, chunk, embed, version, archive, search, gap analytics.
Extends the existing /ai/knowledge endpoint without removing it.
"""

from __future__ import annotations

import csv
import io
from datetime import datetime, timezone
from typing import Annotated

from fastapi import APIRouter, Depends, File, HTTPException, Query, UploadFile
from pydantic import BaseModel, Field
from sqlalchemy import func, select
from sqlalchemy.ext.asyncio import AsyncSession

from database import get_session
from models import AuditEvent, KBGapLog, KnowledgeDocument, User
from security import assert_tenant, require_roles
from services.ai_service import AIConfigurationError, embed_text
from services.vector_store import upsert_vector

router = APIRouter()


class KBDocumentCreate(BaseModel):
    title: str = Field(min_length=1, max_length=512)
    body: str = Field(min_length=1)
    category: str = Field(default="general", max_length=128)
    doc_type: str = Field(default="article", max_length=64)
    source_uri: str | None = Field(default=None, max_length=1024)
    metadata: dict = Field(default_factory=dict)


class KBDocumentUpdate(BaseModel):
    title: str | None = None
    body: str | None = None
    category: str | None = None
    status: str | None = None

    def has_updates(self) -> bool:
        return any(v is not None for v in [self.title, self.body, self.category, self.status])


SUPPORTED_UPLOAD_TYPES = {".pdf", ".docx", ".txt", ".csv", ".md"}
MAX_UPLOAD_BYTES = 8 * 1024 * 1024


def _extension(filename: str) -> str:
    return "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""


def _chunk_text(text: str, *, chunk_size: int = 2200, overlap: int = 250) -> list[str]:
    normalized = "\n".join(line.strip() for line in text.splitlines() if line.strip())
    if len(normalized) <= chunk_size:
        return [normalized] if normalized else []
    chunks: list[str] = []
    start = 0
    while start < len(normalized):
        end = min(start + chunk_size, len(normalized))
        boundary = normalized.rfind("\n", start, end)
        if boundary <= start + 400:
            boundary = end
        chunk = normalized[start:boundary].strip()
        if chunk:
            chunks.append(chunk)
        start = boundary if boundary >= len(normalized) else max(0, boundary - overlap)
    return chunks


def _extract_upload_text(filename: str, content: bytes) -> str:
    ext = _extension(filename)
    if ext in {".txt", ".md"}:
        return content.decode("utf-8", errors="replace")
    if ext == ".csv":
        decoded = content.decode("utf-8", errors="replace")
        rows = csv.reader(io.StringIO(decoded))
        return "\n".join(" | ".join(cell.strip() for cell in row) for row in rows)
    if ext == ".pdf":
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise HTTPException(status_code=500, detail="PDF ingestion requires pypdf") from exc
        reader = PdfReader(io.BytesIO(content))
        return "\n".join((page.extract_text() or "").strip() for page in reader.pages)
    if ext == ".docx":
        try:
            from docx import Document
        except ImportError as exc:
            raise HTTPException(status_code=500, detail="DOCX ingestion requires python-docx") from exc
        document = Document(io.BytesIO(content))
        return "\n".join(paragraph.text for paragraph in document.paragraphs)
    raise HTTPException(status_code=400, detail="Unsupported file type")


def _doc_summary(doc: KnowledgeDocument) -> dict:
    return {
        "id": doc.id,
        "title": doc.title,
        "category": doc.category,
        "doc_type": doc.doc_type,
        "status": doc.status,
        "version": doc.version,
        "source_uri": doc.source_uri,
        "file_type": doc.file_type,
        "file_size_bytes": doc.file_size_bytes,
        "chunk_count": doc.chunk_count,
        "last_indexed_at": doc.last_indexed_at.isoformat() if doc.last_indexed_at else None,
        "created_at": doc.created_at.isoformat(),
        "updated_at": doc.updated_at.isoformat(),
    }


@router.get("/knowledge")
async def list_documents(
    user: Annotated[User, Depends(require_roles("tenant_admin", "manager", "qa_reviewer", "read_only_analyst"))],
    session: Annotated[AsyncSession, Depends(get_session)],
    tenant_id: str | None = None,
    category: str | None = None,
    status: str = "active",
    limit: int = Query(default=100, le=500),
):
    """List all knowledge base documents for the tenant."""
    scoped_tenant = assert_tenant(user, tenant_id)
    query = select(KnowledgeDocument).where(KnowledgeDocument.tenant_id == scoped_tenant)
    if category:
        query = query.where(KnowledgeDocument.category == category)
    if status != "all":
        query = query.where(KnowledgeDocument.status == status)
    query = query.order_by(KnowledgeDocument.updated_at.desc()).limit(limit)

    docs = (await session.scalars(query)).all()
    total = await session.scalar(
        select(func.count(KnowledgeDocument.id)).where(KnowledgeDocument.tenant_id == scoped_tenant)
    )

    # Category breakdown
    categories = (
        await session.execute(
            select(KnowledgeDocument.category, func.count(KnowledgeDocument.id))
            .where(KnowledgeDocument.tenant_id == scoped_tenant, KnowledgeDocument.status == "active")
            .group_by(KnowledgeDocument.category)
        )
    ).all()

    return {
        "documents": [_doc_summary(doc) for doc in docs],
        "total": total or 0,
        "categories": {cat: count for cat, count in categories},
    }


@router.post("/knowledge")
async def create_document(
    body: KBDocumentCreate,
    user: Annotated[User, Depends(require_roles("tenant_admin"))],
    session: Annotated[AsyncSession, Depends(get_session)],
):
    """Create a new knowledge base document with embedding."""
    tenant_id = assert_tenant(user, None)
    try:
        embedding = await embed_text(f"{body.title}\n\n{body.body}")
    except AIConfigurationError as exc:
        raise HTTPException(status_code=503, detail=str(exc)) from exc

    from datetime import datetime, timezone
    doc = KnowledgeDocument(
        tenant_id=tenant_id,
        title=body.title,
        body=body.body,
        category=body.category,
        doc_type=body.doc_type,
        source_uri=body.source_uri,
        source_metadata=body.metadata,
        embedding=embedding,
        created_by=user.id,
        last_indexed_at=datetime.now(timezone.utc),
    )
    session.add(doc)

    session.add(AuditEvent(
        tenant_id=tenant_id,
        user_id=user.id,
        action="kb.create",
        resource_type="knowledge_document",
        resource_id=doc.id,
        details={"title": body.title, "category": body.category},
    ))

    await session.commit()
    await session.refresh(doc)

    # Upsert to vector store
    await upsert_vector(
        tenant_id=tenant_id,
        bucket="knowledge",
        vector_id=doc.id,
        vector=embedding,
        metadata={"title": doc.title, "source_uri": doc.source_uri or "", "category": doc.category, "body": doc.body[:4000]},
    )

    return {"status": "created", "document": _doc_summary(doc)}


@router.post("/knowledge/upload")
async def upload_document(
    user: Annotated[User, Depends(require_roles("tenant_admin"))],
    session: Annotated[AsyncSession, Depends(get_session)],
    file: UploadFile = File(...),
    category: str = "general",
    doc_type: str = "policy",
):
    """Upload and index a PDF, DOCX, TXT, MD, or CSV knowledge document."""
    tenant_id = assert_tenant(user, None)
    filename = file.filename or "knowledge_upload"
    ext = _extension(filename)
    if ext not in SUPPORTED_UPLOAD_TYPES:
        raise HTTPException(status_code=400, detail=f"Unsupported file type. Allowed: {sorted(SUPPORTED_UPLOAD_TYPES)}")
    content = await file.read()
    if not content or len(content) > MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=400, detail="File must be non-empty and at most 8 MB")

    body = _extract_upload_text(filename, content)
    chunks = _chunk_text(body)
    if not chunks:
        raise HTTPException(status_code=400, detail="No extractable text found in uploaded document")

    parent_id: str | None = None
    docs: list[KnowledgeDocument] = []
    now = datetime.now(timezone.utc)
    for index, chunk in enumerate(chunks, start=1):
        try:
            embedding = await embed_text(f"{filename}\n\n{chunk}")
        except AIConfigurationError as exc:
            raise HTTPException(status_code=503, detail=str(exc)) from exc
        doc = KnowledgeDocument(
            tenant_id=tenant_id,
            title=filename if len(chunks) == 1 else f"{filename} chunk {index}",
            body=chunk,
            category=category,
            doc_type=doc_type,
            source_uri=f"upload://{filename}",
            source_metadata={"filename": filename, "chunk_index": index, "chunk_total": len(chunks), "parent_id": parent_id},
            embedding=embedding,
            file_type=ext.lstrip("."),
            file_size_bytes=len(content),
            chunk_count=len(chunks),
            created_by=user.id,
            last_indexed_at=now,
        )
        session.add(doc)
        await session.flush()
        if parent_id is None:
            parent_id = doc.id
            doc.source_metadata = {**doc.source_metadata, "parent_id": parent_id}
        await upsert_vector(
            tenant_id=tenant_id,
            bucket="knowledge",
            vector_id=doc.id,
            vector=embedding,
            metadata={
                "title": doc.title,
                "source_uri": doc.source_uri or "",
                "category": doc.category,
                "body": doc.body[:4000],
                "filename": filename,
                "chunk_index": index,
            },
        )
        docs.append(doc)

    session.add(AuditEvent(
        tenant_id=tenant_id,
        user_id=user.id,
        action="kb.upload",
        resource_type="knowledge_document",
        resource_id=parent_id,
        details={"filename": filename, "chunks": len(chunks), "file_type": ext.lstrip(".")},
    ))
    await session.commit()
    return {"status": "indexed", "documents": [_doc_summary(doc) for doc in docs], "chunk_count": len(docs)}


@router.get("/knowledge/{doc_id}")
async def get_document(
    doc_id: str,
    user: Annotated[User, Depends(require_roles("tenant_admin", "manager", "qa_reviewer", "support_agent", "read_only_analyst"))],
    session: Annotated[AsyncSession, Depends(get_session)],
):
    """Get a single knowledge base document with full body."""
    doc = await session.get(KnowledgeDocument, doc_id)
    if doc is None:
        raise HTTPException(status_code=404, detail="Document not found")
    assert_tenant(user, doc.tenant_id)
    result = _doc_summary(doc)
    result["body"] = doc.body
    result["source_metadata"] = doc.source_metadata
    return result


@router.patch("/knowledge/{doc_id}")
async def update_document(
    doc_id: str,
    body: KBDocumentUpdate,
    user: Annotated[User, Depends(require_roles("tenant_admin"))],
    session: Annotated[AsyncSession, Depends(get_session)],
):
    """Update a knowledge base document. Increments version on body change."""
    doc = await session.get(KnowledgeDocument, doc_id)
    if doc is None:
        raise HTTPException(status_code=404, detail="Document not found")
    assert_tenant(user, doc.tenant_id)

    if not body.has_updates():
        raise HTTPException(status_code=400, detail="No updates provided")

    previous = {"title": doc.title, "status": doc.status, "version": doc.version}
    re_embed = False

    if body.title is not None:
        doc.title = body.title
        re_embed = True
    if body.body is not None:
        doc.body = body.body
        doc.version += 1
        re_embed = True
    if body.category is not None:
        doc.category = body.category
    if body.status is not None:
        doc.status = body.status

    if re_embed:
        try:
            embedding = await embed_text(f"{doc.title}\n\n{doc.body}")
            doc.embedding = embedding
            from datetime import datetime, timezone
            doc.last_indexed_at = datetime.now(timezone.utc)
            await upsert_vector(
                tenant_id=doc.tenant_id,
                bucket="knowledge",
                vector_id=doc.id,
                vector=embedding,
                metadata={"title": doc.title, "source_uri": doc.source_uri or "", "category": doc.category, "body": doc.body[:4000]},
            )
        except AIConfigurationError:
            pass  # Continue even if re-embedding fails

    session.add(AuditEvent(
        tenant_id=doc.tenant_id,
        user_id=user.id,
        action="kb.update",
        resource_type="knowledge_document",
        resource_id=doc_id,
        previous_state=previous,
        new_state={"title": doc.title, "status": doc.status, "version": doc.version},
    ))

    await session.commit()
    return {"status": "updated", "document": _doc_summary(doc)}


@router.delete("/knowledge/{doc_id}")
async def archive_document(
    doc_id: str,
    user: Annotated[User, Depends(require_roles("tenant_admin"))],
    session: Annotated[AsyncSession, Depends(get_session)],
):
    """Archive (soft-delete) a knowledge base document."""
    doc = await session.get(KnowledgeDocument, doc_id)
    if doc is None:
        raise HTTPException(status_code=404, detail="Document not found")
    assert_tenant(user, doc.tenant_id)
    doc.status = "archived"

    session.add(AuditEvent(
        tenant_id=doc.tenant_id,
        user_id=user.id,
        action="kb.archive",
        resource_type="knowledge_document",
        resource_id=doc_id,
    ))

    await session.commit()
    return {"status": "archived", "id": doc_id}


# ── KB Gap Analytics ─────────────────────────────────────────

@router.get("/knowledge/gaps/analytics")
async def get_kb_gaps(
    user: Annotated[User, Depends(require_roles("tenant_admin", "manager", "qa_reviewer", "read_only_analyst"))],
    session: Annotated[AsyncSession, Depends(get_session)],
    tenant_id: str | None = None,
    resolved: bool | None = None,
):
    """Get KB gap analytics — queries where AI confidence was low."""
    scoped_tenant = assert_tenant(user, tenant_id)
    query = select(KBGapLog).where(KBGapLog.tenant_id == scoped_tenant)
    if resolved is not None:
        query = query.where(KBGapLog.resolved == resolved)
    gaps = (await session.scalars(query.order_by(KBGapLog.created_at.desc()).limit(200))).all()

    return {
        "gaps": [
            {
                "id": gap.id,
                "query": gap.query,
                "ai_confidence": gap.ai_confidence,
                "suggested_topic": gap.suggested_topic,
                "resolved": gap.resolved,
                "ticket_id": gap.ticket_id,
                "created_at": gap.created_at.isoformat(),
            }
            for gap in gaps
        ],
        "total": len(gaps),
        "unresolved_count": sum(1 for g in gaps if not g.resolved),
    }


@router.patch("/knowledge/gaps/{gap_id}/resolve")
async def resolve_gap(
    gap_id: str,
    user: Annotated[User, Depends(require_roles("tenant_admin"))],
    session: Annotated[AsyncSession, Depends(get_session)],
):
    """Mark a KB gap as resolved."""
    gap = await session.get(KBGapLog, gap_id)
    if gap is None:
        raise HTTPException(status_code=404, detail="Gap not found")
    assert_tenant(user, gap.tenant_id)
    gap.resolved = True
    await session.commit()
    return {"status": "resolved", "id": gap_id}
